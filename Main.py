import requests
import pandas as pd
from docx import Document
from openpyxl.worksheet.pagebreak import Break


# Function to call the OpenAI ChatGPT API
def call_chatgpt(prompt, openai_api_key):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {openai_api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "gpt-4o",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 4096  # Adjust as needed
    }
    response = requests.post(url, headers=headers, json=data)

    # Check for errors in the response
    if response.status_code == 200:
        response_json = response.json()
        if 'choices' in response_json and len(response_json['choices']) > 0:
            return response_json['choices'][0]['message']['content']
        else:
            return "Error: No choices found in the API response."
    else:
        return f"Error: API request failed with status code {response.status_code}. Response: {response.text}"


# Function to write results to a Word document
def write_results_to_docx(filename, data):
    doc = Document()
    doc.add_heading('ChatGPT - Essay Section', level=1)

    # Loop through data (list of tuples) and write each question/response to the document
    for i, (prompt, response) in enumerate(data):
        doc.add_heading(f"Question {i + 1}: {prompt}", level=2)
        doc.add_paragraph(f"ChatGPT Response: {response}")
        #doc.add_paragraph("\n")  # Blank line between entries

    doc.save(filename)


# Main function to read prompts from Excel and store responses in a Word document
def main():
    openai_api_key = ""  # Replace with your OpenAI API key
    input_excel_file = r"C:\Users\VINEET_NEW\OneDrive\Desktop\IAS\Essay - Prompts.xlsx"  # Replace with your Excel file containing prompts
    output_file = r"C:\Users\VINEET_NEW\OneDrive\Desktop\IAS\Responses\Topic Info\Essay Topics - ChatGPT.docx"

    # Read Excel file
    df = pd.read_excel(input_excel_file)

    # Assuming the prompts are in a column named 'Prompt'
    if 'Topic' not in df.columns:
        print("Error: Excel file does not contain a 'Prompt' column.")
        return

    prompts = df['Topic'].tolist()

    # Store results (prompt, response) in a list
    results = []

    # Make API calls for each prompt
    for topic in prompts:
        print(f"Processing topic: {topic}")
        prompt1 = f"Give essay topics on {topic} for TSPSC Group 1 main exam"
        print(f"Processing Prompt: {prompt1}")
        response = call_chatgpt(prompt1, openai_api_key)
        results.append((prompt1, response))  # Store prompt and response

        # prompt2 = f"For TSPSC Group 1 main exam, write me one example essay on one of the topics on {topic}"
        # print(f"Processing Prompt: {prompt2}")
        # response = call_chatgpt(prompt2, openai_api_key)
        # results.append((prompt2, response))  # Store prompt and 2nd response

    # Write all results to a Word document
    write_results_to_docx(output_file, results)

    print(f"Responses saved in {output_file}")

if __name__ == "__main__":
    main()
